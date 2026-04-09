"""
Professional Indian Construction Contractor Bill Generator
Generates both PDF (Bill_Final.pdf) and DOCX (Bill_Final.docx)
"""

import os
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import PageBreak

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# BILL DATA
# ─────────────────────────────────────────────

BILL_INFO = {
    "bill_no": "01",
    "date": "09/04/2026",
    "contractor": "M/s. Shri Construction Co.",
    "owner": "Mr. Ramesh Kumar",
    "site": "Plot No. 12, Sector-4, New Colony, Indore (M.P.)",
}

# Section A
CIVIL_WORK = [
    ("1", "Earth Work Excavation for Foundation", "150.00", "Cft", "8.00", "1,200.00"),
    ("2", "PCC Work (1:4:8) for Foundation", "80.00", "Cft", "45.00", "3,600.00"),
    ("3", "RCC Work (1:2:4) for Column, Beam & Slab", "250.00", "Cft", "90.00", "22,500.00"),
    ("4", "Brick Work (9\" Wall)", "500.00", "Cft", "32.00", "16,000.00"),
    ("5", "Brick Work (4.5\" Wall)", "300.00", "Cft", "22.00", "6,600.00"),
    ("6", "Plaster Work (Inside)", "1200.00", "Sft", "12.00", "14,400.00"),
    ("7", "Plaster Work (Outside)", "800.00", "Sft", "14.00", "11,200.00"),
]
CIVIL_SUBTOTAL = "75,500.00"

# Section B
FLOORING_WORK = [
    ("1", "Vitrified Tile Flooring (2x2)", "600.00", "Sft", "28.00", "16,800.00"),
    ("2", "Bathroom Tile Work", "200.00", "Sft", "32.00", "6,400.00"),
    ("3", "Kitchen Tile Work (Wall)", "150.00", "Sft", "30.00", "4,500.00"),
    ("4", "Kota Stone Flooring (Porch Area)", "120.00", "Sft", "25.00", "3,000.00"),
]
FLOORING_SUBTOTAL = "30,700.00"

# Section C - Granite Stone Work (original rows + new skirting sub-items)
GRANITE_WORK_MAIN = [
    ("1", "Granite Counter Top (Kitchen)", "30.00", "Sft", "120.00", "3,600.00"),
    ("2", "Granite Window Sill", "40.00", "Rft", "85.00", "3,400.00"),
    ("3", "Granite Staircase Work", "56.00", "Sft", "110.00", "6,160.00"),
    ("4", "Granite Threshold/Door Frame", "8.00", "Nos", "455.00", "3,640.00"),
]
# Skirting sub-items (measurement details)
GRANITE_SKIRTING_SUB = [
    ("(a)", "Middliya Room Skirting - 72'-0\"", "72.00", "Rft", "", ""),
    ("(b)", "Porch Pillar Skirting (5'-10\" \u00d7 7 Nos.)", "40.81", "Rft", "", ""),
    ("(c)", "VIP Seena", "7.00", "Rft", "", ""),
    ("(d)", "Granite Dell", "4.75", "Rft", "", ""),
    ("(e)", "Kota Dell (4'0\" \u00d7 2 Nos.)", "8.00", "Rft", "", ""),
    ("(f)", "Fire Seena Skirting Extra", "50.00", "Rft", "", ""),
]
SKIRTING_TOTAL_QTY = "182.00"
SKIRTING_RATE = "40.00"
SKIRTING_AMOUNT = "7,280.00"
GRANITE_SUBTOTAL = "24,080.00"

# Section D
PAINTING_WORK = [
    ("1", "Primer Coat (Inside)", "1200.00", "Sft", "4.00", "4,800.00"),
    ("2", "Putty Work (2 Coat)", "1200.00", "Sft", "8.00", "9,600.00"),
    ("3", "Emulsion Paint (Inside - 2 Coat)", "1200.00", "Sft", "7.00", "8,400.00"),
    ("4", "Exterior Paint (Outside)", "800.00", "Sft", "9.00", "7,200.00"),
]
PAINTING_SUBTOTAL = "30,000.00"

# Section E
ELECTRICAL_WORK = [
    ("1", "Wiring Work (Concealed)", "25.00", "Points", "350.00", "8,750.00"),
    ("2", "Main Board with MCB", "1.00", "Nos", "3,500.00", "3,500.00"),
    ("3", "Earthing Work", "2.00", "Nos", "1,500.00", "3,000.00"),
]
ELECTRICAL_SUBTOTAL = "15,250.00"

# Section F
PLUMBING_WORK = [
    ("1", "Water Supply Line (CPVC)", "120.00", "Rft", "35.00", "4,200.00"),
    ("2", "Drainage Line (PVC)", "80.00", "Rft", "28.00", "2,240.00"),
    ("3", "Sanitary Fitting Work", "1.00", "LS", "5,000.00", "5,000.00"),
]
PLUMBING_SUBTOTAL = "11,440.00"

GRAND_TOTAL = "1,86,970.00"
AMOUNT_IN_WORDS = "Rupees One Lakh Eighty-Six Thousand Nine Hundred Seventy Only"


# ─────────────────────────────────────────────
# PDF GENERATION
# ─────────────────────────────────────────────

def fmt_inr(val):
    """Return value prefixed with rupee symbol."""
    return "\u20b9" + val


def build_pdf():
    output_path = os.path.join(os.path.dirname(__file__), "Bill_Final.pdf")
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "BillTitle",
        fontName="Helvetica-Bold",
        fontSize=20,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    center_bold = ParagraphStyle(
        "CenterBold",
        fontName="Helvetica-Bold",
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    normal_center = ParagraphStyle(
        "NormalCenter",
        fontName="Helvetica",
        fontSize=9,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    info_label = ParagraphStyle(
        "InfoLabel",
        fontName="Helvetica-Bold",
        fontSize=9,
        alignment=TA_LEFT,
    )
    info_val = ParagraphStyle(
        "InfoVal",
        fontName="Helvetica",
        fontSize=9,
        alignment=TA_LEFT,
    )
    section_hdr = ParagraphStyle(
        "SectionHdr",
        fontName="Helvetica-Bold",
        fontSize=9,
        alignment=TA_LEFT,
    )
    cell_normal = ParagraphStyle(
        "CellNormal",
        fontName="Helvetica",
        fontSize=8,
        alignment=TA_LEFT,
        leading=11,
    )
    cell_right = ParagraphStyle(
        "CellRight",
        fontName="Helvetica",
        fontSize=8,
        alignment=TA_RIGHT,
        leading=11,
    )
    cell_bold = ParagraphStyle(
        "CellBold",
        fontName="Helvetica-Bold",
        fontSize=8,
        alignment=TA_LEFT,
        leading=11,
    )
    cell_bold_right = ParagraphStyle(
        "CellBoldRight",
        fontName="Helvetica-Bold",
        fontSize=8,
        alignment=TA_RIGHT,
        leading=11,
    )
    cell_center = ParagraphStyle(
        "CellCenter",
        fontName="Helvetica",
        fontSize=8,
        alignment=TA_CENTER,
        leading=11,
    )
    cell_bold_center = ParagraphStyle(
        "CellBoldCenter",
        fontName="Helvetica-Bold",
        fontSize=8,
        alignment=TA_CENTER,
        leading=11,
    )
    indent_style = ParagraphStyle(
        "IndentStyle",
        fontName="Helvetica",
        fontSize=8,
        alignment=TA_LEFT,
        leading=11,
        leftIndent=8,
    )

    # ── Colours ──
    HEADER_BG = colors.HexColor("#D3D3D3")   # light grey for section headers
    COL_HDR_BG = colors.HexColor("#4F81BD")  # blue for column headers
    SUBTOTAL_BG = colors.HexColor("#E8E8E8")
    GRAND_BG = colors.HexColor("#C0C0C0")
    WHITE = colors.white
    BLACK = colors.black

    # Column widths (S.No | Description | Qty | Unit | Rate | Amount)
    col_w = [0.45 * inch, 3.2 * inch, 0.75 * inch, 0.55 * inch, 0.75 * inch, 0.85 * inch]

    def col_header_row():
        hdrs = ["S.No.", "Description of Work", "Quantity", "Unit", "Rate (\u20b9)", "Amount (\u20b9)"]
        return [Paragraph(h, ParagraphStyle("CH", fontName="Helvetica-Bold", fontSize=8,
                                            alignment=TA_CENTER, textColor=WHITE, leading=11))
                for h in hdrs]

    def section_row(label):
        return [
            Paragraph("", cell_bold),
            Paragraph(label, section_hdr),
            "", "", "", "",
        ]

    def data_row(sno, desc, qty, unit, rate, amt, indent=False):
        d_style = indent_style if indent else cell_normal
        return [
            Paragraph(sno, cell_center),
            Paragraph(desc, d_style),
            Paragraph(qty, cell_right),
            Paragraph(unit, cell_center),
            Paragraph(rate, cell_right),
            Paragraph(amt, cell_right),
        ]

    def subtotal_row(label, amount):
        return [
            Paragraph("", cell_bold),
            Paragraph(label, cell_bold),
            "", "", "",
            Paragraph(fmt_inr(amount), cell_bold_right),
        ]

    def build_table(rows, extra_styles=None):
        tbl = Table(rows, colWidths=col_w, repeatRows=1)
        base_style = [
            ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
            ("BACKGROUND", (0, 0), (-1, 0), COL_HDR_BG),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]
        if extra_styles:
            base_style.extend(extra_styles)
        tbl.setStyle(TableStyle(base_style))
        return tbl

    story = []

    # ── PAGE 1 ──────────────────────────────────────────────────────────────

    # Title
    story.append(Paragraph("BILL", title_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLACK))
    story.append(Spacer(1, 4))

    # Bill meta info table
    meta_data = [
        [Paragraph("Bill No:", info_label), Paragraph(BILL_INFO["bill_no"], info_val),
         Paragraph("Date:", info_label), Paragraph(BILL_INFO["date"], info_val)],
        [Paragraph("Contractor:", info_label), Paragraph(BILL_INFO["contractor"], info_val),
         "", ""],
        [Paragraph("Owner/Client:", info_label), Paragraph(BILL_INFO["owner"], info_val),
         "", ""],
        [Paragraph("Site Address:", info_label), Paragraph(BILL_INFO["site"], info_val),
         "", ""],
    ]
    meta_tbl = Table(meta_data, colWidths=[1.0 * inch, 2.7 * inch, 0.85 * inch, 1.8 * inch])
    meta_tbl.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.8, BLACK),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("SPAN", (1, 1), (3, 1)),
        ("SPAN", (1, 2), (3, 2)),
        ("SPAN", (1, 3), (3, 3)),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 8))

    # ── Section A ──
    rows_a = [col_header_row()]
    rows_a.append(section_row("Section A : CIVIL WORK"))
    for r in CIVIL_WORK:
        rows_a.append(data_row(*r))
    rows_a.append(subtotal_row("Sub Total - Civil Work", CIVIL_SUBTOTAL))

    extra_a = [
        ("BACKGROUND", (0, 1), (-1, 1), HEADER_BG),
        ("SPAN", (1, 1), (5, 1)),
        ("BACKGROUND", (0, len(rows_a) - 1), (-1, len(rows_a) - 1), SUBTOTAL_BG),
        ("SPAN", (1, len(rows_a) - 1), (4, len(rows_a) - 1)),
    ]
    story.append(build_table(rows_a, extra_a))
    story.append(Spacer(1, 8))

    # ── Section B ──
    rows_b = [col_header_row()]
    rows_b.append(section_row("Section B : FLOORING WORK"))
    for r in FLOORING_WORK:
        rows_b.append(data_row(*r))
    rows_b.append(subtotal_row("Sub Total - Flooring Work", FLOORING_SUBTOTAL))

    extra_b = [
        ("BACKGROUND", (0, 1), (-1, 1), HEADER_BG),
        ("SPAN", (1, 1), (5, 1)),
        ("BACKGROUND", (0, len(rows_b) - 1), (-1, len(rows_b) - 1), SUBTOTAL_BG),
        ("SPAN", (1, len(rows_b) - 1), (4, len(rows_b) - 1)),
    ]
    story.append(build_table(rows_b, extra_b))

    # ── Page break ──
    story.append(PageBreak())

    # ── PAGE 2 ──────────────────────────────────────────────────────────────

    # ── Section C (Granite Stone Work) ──
    rows_c = [col_header_row()]
    rows_c.append(section_row("Section C : GRANITE STONE WORK"))
    for r in GRANITE_WORK_MAIN:
        rows_c.append(data_row(*r))

    # Row 5 header "Granite Stone Skirting"
    rows_c.append(data_row("5", "Granite Stone Skirting", "", "", "", ""))

    # Measurement details label row
    meas_row_idx = len(rows_c)
    rows_c.append([
        Paragraph("", cell_normal),
        Paragraph("Measurement Details:", ParagraphStyle("italic8", fontName="Helvetica-Oblique",
                                                         fontSize=8, alignment=TA_LEFT, leading=11)),
        "", "", "", "",
    ])

    # Sub-items
    for sub in GRANITE_SKIRTING_SUB:
        rows_c.append(data_row(sub[0], sub[1], sub[2], sub[3], sub[4], sub[5], indent=True))

    # Skirting total row
    skirting_total_idx = len(rows_c)
    rows_c.append([
        Paragraph("", cell_bold),
        Paragraph("Total Granite Stone Skirting", cell_bold),
        Paragraph(SKIRTING_TOTAL_QTY, cell_bold_right),
        Paragraph("Rft", cell_bold_center),
        Paragraph(SKIRTING_RATE, cell_bold_right),
        Paragraph(fmt_inr(SKIRTING_AMOUNT), cell_bold_right),
    ])

    rows_c.append(subtotal_row("Sub Total - Granite Stone Work", GRANITE_SUBTOTAL))

    extra_c = [
        ("BACKGROUND", (0, 1), (-1, 1), HEADER_BG),
        ("SPAN", (1, 1), (5, 1)),
        # meas details span
        ("SPAN", (1, meas_row_idx), (5, meas_row_idx)),
        # skirting total
        ("BACKGROUND", (0, skirting_total_idx), (-1, skirting_total_idx), SUBTOTAL_BG),
        # subtotal
        ("BACKGROUND", (0, len(rows_c) - 1), (-1, len(rows_c) - 1), SUBTOTAL_BG),
        ("SPAN", (1, len(rows_c) - 1), (4, len(rows_c) - 1)),
    ]
    story.append(build_table(rows_c, extra_c))
    story.append(Spacer(1, 8))

    # ── Section D ──
    rows_d = [col_header_row()]
    rows_d.append(section_row("Section D : PAINTING WORK"))
    for r in PAINTING_WORK:
        rows_d.append(data_row(*r))
    rows_d.append(subtotal_row("Sub Total - Painting Work", PAINTING_SUBTOTAL))

    extra_d = [
        ("BACKGROUND", (0, 1), (-1, 1), HEADER_BG),
        ("SPAN", (1, 1), (5, 1)),
        ("BACKGROUND", (0, len(rows_d) - 1), (-1, len(rows_d) - 1), SUBTOTAL_BG),
        ("SPAN", (1, len(rows_d) - 1), (4, len(rows_d) - 1)),
    ]
    story.append(build_table(rows_d, extra_d))

    # ── Page break ──
    story.append(PageBreak())

    # ── PAGE 3 ──────────────────────────────────────────────────────────────

    # ── Section E ──
    rows_e = [col_header_row()]
    rows_e.append(section_row("Section E : ELECTRICAL WORK"))
    for r in ELECTRICAL_WORK:
        rows_e.append(data_row(*r))
    rows_e.append(subtotal_row("Sub Total - Electrical Work", ELECTRICAL_SUBTOTAL))

    extra_e = [
        ("BACKGROUND", (0, 1), (-1, 1), HEADER_BG),
        ("SPAN", (1, 1), (5, 1)),
        ("BACKGROUND", (0, len(rows_e) - 1), (-1, len(rows_e) - 1), SUBTOTAL_BG),
        ("SPAN", (1, len(rows_e) - 1), (4, len(rows_e) - 1)),
    ]
    story.append(build_table(rows_e, extra_e))
    story.append(Spacer(1, 8))

    # ── Section F ──
    rows_f = [col_header_row()]
    rows_f.append(section_row("Section F : PLUMBING WORK"))
    for r in PLUMBING_WORK:
        rows_f.append(data_row(*r))
    rows_f.append(subtotal_row("Sub Total - Plumbing Work", PLUMBING_SUBTOTAL))

    extra_f = [
        ("BACKGROUND", (0, 1), (-1, 1), HEADER_BG),
        ("SPAN", (1, 1), (5, 1)),
        ("BACKGROUND", (0, len(rows_f) - 1), (-1, len(rows_f) - 1), SUBTOTAL_BG),
        ("SPAN", (1, len(rows_f) - 1), (4, len(rows_f) - 1)),
    ]
    story.append(build_table(rows_f, extra_f))
    story.append(Spacer(1, 12))

    # ── Grand Total Summary ──
    story.append(HRFlowable(width="100%", thickness=1.0, color=BLACK))
    story.append(Spacer(1, 4))

    gt_hdr_style = ParagraphStyle("GTHdr", fontName="Helvetica-Bold", fontSize=10,
                                  alignment=TA_CENTER, textColor=WHITE)
    gt_label = ParagraphStyle("GTLabel", fontName="Helvetica-Bold", fontSize=9, alignment=TA_LEFT)
    gt_val = ParagraphStyle("GTVal", fontName="Helvetica", fontSize=9, alignment=TA_RIGHT)
    gt_val_bold = ParagraphStyle("GTValB", fontName="Helvetica-Bold", fontSize=10, alignment=TA_RIGHT)
    gt_label_grand = ParagraphStyle("GTLabelG", fontName="Helvetica-Bold", fontSize=10,
                                    alignment=TA_LEFT)

    summary_rows = [
        [Paragraph("GRAND TOTAL SUMMARY", gt_hdr_style), ""],
        [Paragraph("A.  Civil Work", gt_label), Paragraph(fmt_inr("75,500.00"), gt_val)],
        [Paragraph("B.  Flooring Work", gt_label), Paragraph(fmt_inr("30,700.00"), gt_val)],
        [Paragraph("C.  Granite Stone Work", gt_label), Paragraph(fmt_inr("24,080.00"), gt_val)],
        [Paragraph("D.  Painting Work", gt_label), Paragraph(fmt_inr("30,000.00"), gt_val)],
        [Paragraph("E.  Electrical Work", gt_label), Paragraph(fmt_inr("15,250.00"), gt_val)],
        [Paragraph("F.  Plumbing Work", gt_label), Paragraph(fmt_inr("11,440.00"), gt_val)],
        [Paragraph("GRAND TOTAL", gt_label_grand), Paragraph(fmt_inr(GRAND_TOTAL), gt_val_bold)],
    ]

    gt_col_w = [4.0 * inch, 2.55 * inch]
    gt_tbl = Table(summary_rows, colWidths=gt_col_w)
    gt_tbl.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 1.0, BLACK),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, BLACK),
        ("SPAN", (0, 0), (1, 0)),
        ("BACKGROUND", (0, 0), (-1, 0), COL_HDR_BG),
        ("BACKGROUND", (0, 7), (-1, 7), GRAND_BG),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(gt_tbl)
    story.append(Spacer(1, 6))

    # Amount in words
    words_style = ParagraphStyle("Words", fontName="Helvetica-Bold", fontSize=9,
                                 alignment=TA_CENTER)
    story.append(Paragraph(
        f"Amount in Words : <b>{AMOUNT_IN_WORDS}</b>",
        ParagraphStyle("Words2", fontName="Helvetica", fontSize=9, alignment=TA_CENTER)
    ))
    story.append(Spacer(1, 24))

    # ── Signature Block ──
    sig_data = [
        [Paragraph("Contractor Signature", ParagraphStyle("sig", fontName="Helvetica-Bold",
                                                          fontSize=9, alignment=TA_CENTER)),
         Paragraph("Owner / Client Signature", ParagraphStyle("sig2", fontName="Helvetica-Bold",
                                                               fontSize=9, alignment=TA_CENTER))],
        [Paragraph("_______________________________",
                   ParagraphStyle("sigline", fontName="Helvetica", fontSize=9, alignment=TA_CENTER)),
         Paragraph("_______________________________",
                   ParagraphStyle("sigline2", fontName="Helvetica", fontSize=9, alignment=TA_CENTER))],
        [Paragraph(BILL_INFO["contractor"],
                   ParagraphStyle("signame", fontName="Helvetica", fontSize=8, alignment=TA_CENTER)),
         Paragraph(BILL_INFO["owner"],
                   ParagraphStyle("signame2", fontName="Helvetica", fontSize=8, alignment=TA_CENTER))],
    ]
    sig_tbl = Table(sig_data, colWidths=[3.28 * inch, 3.28 * inch])
    sig_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("BOX", (0, 0), (0, -1), 0.5, BLACK),
        ("BOX", (1, 0), (1, -1), 0.5, BLACK),
    ]))
    story.append(sig_tbl)

    # ── Page number footer via canvas callback ──
    def add_page_number(canvas_obj, doc_obj):
        canvas_obj.saveState()
        canvas_obj.setFont("Helvetica", 8)
        page_num_text = f"Page {canvas_obj.getPageNumber()}"
        canvas_obj.drawCentredString(306, 20, page_num_text)
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    print(f"PDF generated: {output_path}")


# ─────────────────────────────────────────────
# DOCX GENERATION
# ─────────────────────────────────────────────

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Add borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), val.get("val", "single"))
            border.set(qn("w:sz"), str(val.get("sz", 4)))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill_color):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def apply_all_borders(table, sz=4, color="000000"):
    """Apply borders to all cells of a table."""
    border_val = {"val": "single", "sz": sz, "color": color}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_val, bottom=border_val,
                            left=border_val, right=border_val)


def add_paragraph_in_cell(cell, text, bold=False, italic=False,
                           font_name="Calibri", font_size=9,
                           align=WD_ALIGN_PARAGRAPH.LEFT, color_hex=None):
    """Clear cell and add a single paragraph."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    # Reduce cell spacing
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


def set_col_widths(table, widths_cm):
    """Set column widths in a table."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(widths_cm[i] * 567)))  # 567 twips per cm
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


COL_HDR_HEX = "4F81BD"
SECTION_HDR_HEX = "D3D3D3"
SUBTOTAL_HEX = "E8E8E8"
GRAND_HEX = "C0C0C0"

# Column widths in cm
COL_WIDTHS = [1.0, 7.5, 1.7, 1.2, 1.7, 2.0]


def add_col_headers(table):
    """Add column header row to a work table."""
    row = table.add_row()
    hdrs = ["S.No.", "Description of Work", "Quantity", "Unit", "Rate (\u20b9)", "Amount (\u20b9)"]
    aligns = [WD_ALIGN_PARAGRAPH.CENTER] * 6
    for i, (h, al) in enumerate(zip(hdrs, aligns)):
        add_paragraph_in_cell(row.cells[i], h, bold=True, font_size=9,
                               align=al, color_hex="FFFFFF")
        set_cell_shading(row.cells[i], COL_HDR_HEX)


def add_section_header_row(table, label):
    """Add a merged section header row."""
    row = table.add_row()
    add_paragraph_in_cell(row.cells[0], "", bold=True, font_size=9)
    add_paragraph_in_cell(row.cells[1], label, bold=True, font_size=9,
                           align=WD_ALIGN_PARAGRAPH.LEFT)
    for i in range(2, 6):
        add_paragraph_in_cell(row.cells[i], "", font_size=9)
    # Merge cells 1-5
    row.cells[1].merge(row.cells[5])
    set_cell_shading(row.cells[0], SECTION_HDR_HEX)
    set_cell_shading(row.cells[1], SECTION_HDR_HEX)


def add_data_row_docx(table, sno, desc, qty, unit, rate, amt,
                       bold=False, indent=False, shading_hex=None):
    """Add a data row to a work table."""
    row = table.add_row()
    vals = [sno, desc, qty, unit, rate, amt]
    aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
    ]
    for i, (v, al) in enumerate(zip(vals, aligns)):
        text = v
        if i in (4, 5) and v and v not in ("", "—"):
            if not v.startswith("\u20b9"):
                text = v
        add_paragraph_in_cell(row.cells[i], text, bold=bold, font_size=9,
                               align=al,
                               italic=(indent and i == 1))
        if shading_hex:
            set_cell_shading(row.cells[i], shading_hex)


def add_subtotal_row_docx(table, label, amount):
    """Add a subtotal row spanning columns 1-4 with amount in col 5."""
    row = table.add_row()
    add_paragraph_in_cell(row.cells[0], "", bold=True, font_size=9)
    add_paragraph_in_cell(row.cells[1], label, bold=True, font_size=9)
    for i in range(2, 5):
        add_paragraph_in_cell(row.cells[i], "", font_size=9)
    add_paragraph_in_cell(row.cells[5], fmt_inr(amount), bold=True, font_size=9,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)
    row.cells[1].merge(row.cells[4])
    for cell in row.cells:
        set_cell_shading(cell, SUBTOTAL_HEX)


def make_work_table(doc, section_title, data_rows, subtotal_label, subtotal_amt,
                     extra_rows_fn=None):
    """Create a full work section table."""
    table = doc.add_table(rows=0, cols=6)
    table.style = "Table Grid"
    add_col_headers(table)
    add_section_header_row(table, section_title)
    for r in data_rows:
        add_data_row_docx(table, *r)
    if extra_rows_fn:
        extra_rows_fn(table)
    add_subtotal_row_docx(table, subtotal_label, subtotal_amt)
    apply_all_borders(table)
    set_col_widths(table, COL_WIDTHS)
    return table


def build_docx():
    output_path = os.path.join(os.path.dirname(__file__), "Bill_Final.docx")
    doc = Document()

    # Page setup - Letter size with margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("CONSTRUCTION CONTRACTOR BILL")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    # Footer with page number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_xml = (
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:instr=" PAGE "><w:r><w:rPr><w:sz w:val="16"/></w:rPr>'
        "<w:t>1</w:t></w:r></w:fldSimple>"
    )
    footer_para._p.append(OxmlElement("w:fldSimple"))
    footer_para.clear()
    from lxml import etree
    fld_el = etree.fromstring(fld_xml)
    footer_para._p.append(fld_el)

    # ── Title ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BILL")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = "Calibri"
    pPr = title_p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "60")
    pPr.append(spacing)

    # ── Bill Meta Info Table ──
    meta_tbl = doc.add_table(rows=4, cols=4)
    meta_tbl.style = "Table Grid"
    meta_data = [
        ("Bill No:", BILL_INFO["bill_no"], "Date:", BILL_INFO["date"]),
        ("Contractor:", BILL_INFO["contractor"], "", ""),
        ("Owner/Client:", BILL_INFO["owner"], "", ""),
        ("Site Address:", BILL_INFO["site"], "", ""),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(meta_data):
        row = meta_tbl.rows[row_idx]
        add_paragraph_in_cell(row.cells[0], l1, bold=True, font_size=9)
        add_paragraph_in_cell(row.cells[1], v1, font_size=9)
        add_paragraph_in_cell(row.cells[2], l2, bold=True, font_size=9)
        add_paragraph_in_cell(row.cells[3], v2, font_size=9)

    # Merge cells for rows where label spans
    for ri in range(1, 4):
        meta_tbl.rows[ri].cells[1].merge(meta_tbl.rows[ri].cells[3])

    apply_all_borders(meta_tbl)
    meta_col_widths = [2.2, 6.0, 1.8, 4.0]
    set_col_widths(meta_tbl, meta_col_widths)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section A: Civil Work ──
    make_work_table(doc, "Section A : CIVIL WORK", CIVIL_WORK,
                    "Sub Total - Civil Work", CIVIL_SUBTOTAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section B: Flooring Work ──
    make_work_table(doc, "Section B : FLOORING WORK", FLOORING_WORK,
                    "Sub Total - Flooring Work", FLOORING_SUBTOTAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section C: Granite Stone Work (with sub-items) ──
    def granite_extra_rows(table):
        # Row 5 - "Granite Stone Skirting" header
        add_data_row_docx(table, "5", "Granite Stone Skirting", "", "", "", "", bold=True)
        # Measurement details label
        mrow = table.add_row()
        add_paragraph_in_cell(mrow.cells[0], "", font_size=9)
        add_paragraph_in_cell(mrow.cells[1], "Measurement Details:", italic=True,
                               font_size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        for i in range(2, 6):
            add_paragraph_in_cell(mrow.cells[i], "", font_size=9)
        mrow.cells[1].merge(mrow.cells[5])
        # Sub-items
        for sub in GRANITE_SKIRTING_SUB:
            add_data_row_docx(table, sub[0], sub[1], sub[2], sub[3], sub[4], sub[5], indent=True)
        # Skirting total
        skt_row = table.add_row()
        add_paragraph_in_cell(skt_row.cells[0], "", bold=True, font_size=9)
        add_paragraph_in_cell(skt_row.cells[1], "Total Granite Stone Skirting",
                               bold=True, font_size=9)
        add_paragraph_in_cell(skt_row.cells[2], SKIRTING_TOTAL_QTY, bold=True,
                               font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_paragraph_in_cell(skt_row.cells[3], "Rft", bold=True, font_size=9,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph_in_cell(skt_row.cells[4], SKIRTING_RATE, bold=True,
                               font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_paragraph_in_cell(skt_row.cells[5], fmt_inr(SKIRTING_AMOUNT), bold=True,
                               font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for cell in skt_row.cells:
            set_cell_shading(cell, SUBTOTAL_HEX)

    make_work_table(doc, "Section C : GRANITE STONE WORK", GRANITE_WORK_MAIN,
                    "Sub Total - Granite Stone Work", GRANITE_SUBTOTAL,
                    extra_rows_fn=granite_extra_rows)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section D: Painting Work ──
    make_work_table(doc, "Section D : PAINTING WORK", PAINTING_WORK,
                    "Sub Total - Painting Work", PAINTING_SUBTOTAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section E: Electrical Work ──
    make_work_table(doc, "Section E : ELECTRICAL WORK", ELECTRICAL_WORK,
                    "Sub Total - Electrical Work", ELECTRICAL_SUBTOTAL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section F: Plumbing Work ──
    make_work_table(doc, "Section F : PLUMBING WORK", PLUMBING_WORK,
                    "Sub Total - Plumbing Work", PLUMBING_SUBTOTAL)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ── Grand Total Summary ──
    gt_hdr_p = doc.add_paragraph()
    gt_hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gt_run = gt_hdr_p.add_run("GRAND TOTAL SUMMARY")
    gt_run.bold = True
    gt_run.font.size = Pt(11)
    gt_run.font.name = "Calibri"

    gt_tbl = doc.add_table(rows=0, cols=2)
    gt_tbl.style = "Table Grid"
    gt_sections = [
        ("A.  Civil Work", "\u20b975,500.00"),
        ("B.  Flooring Work", "\u20b930,700.00"),
        ("C.  Granite Stone Work", "\u20b924,080.00"),
        ("D.  Painting Work", "\u20b930,000.00"),
        ("E.  Electrical Work", "\u20b915,250.00"),
        ("F.  Plumbing Work", "\u20b911,440.00"),
    ]
    for label, amt in gt_sections:
        row = gt_tbl.add_row()
        add_paragraph_in_cell(row.cells[0], label, font_size=9)
        add_paragraph_in_cell(row.cells[1], amt, font_size=9,
                               align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Grand total row
    grand_row = gt_tbl.add_row()
    add_paragraph_in_cell(grand_row.cells[0], "GRAND TOTAL", bold=True, font_size=11)
    add_paragraph_in_cell(grand_row.cells[1], fmt_inr(GRAND_TOTAL), bold=True,
                           font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in grand_row.cells:
        set_cell_shading(cell, GRAND_HEX)

    apply_all_borders(gt_tbl)
    set_col_widths(gt_tbl, [9.5, 4.5])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Amount in words
    words_p = doc.add_paragraph()
    words_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wr = words_p.add_run(f"Amount in Words : {AMOUNT_IN_WORDS}")
    wr.bold = True
    wr.font.size = Pt(9)
    wr.font.name = "Calibri"

    doc.add_paragraph()

    # ── Signature Block ──
    sig_tbl = doc.add_table(rows=3, cols=2)
    sig_tbl.style = "Table Grid"
    sig_data_rows = [
        ("Contractor Signature", "Owner / Client Signature"),
        ("_______________________________", "_______________________________"),
        (BILL_INFO["contractor"], BILL_INFO["owner"]),
    ]
    for ri, (l, r) in enumerate(sig_data_rows):
        row = sig_tbl.rows[ri]
        add_paragraph_in_cell(row.cells[0], l, bold=(ri == 0), font_size=9,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph_in_cell(row.cells[1], r, bold=(ri == 0), font_size=9,
                               align=WD_ALIGN_PARAGRAPH.CENTER)

    apply_all_borders(sig_tbl)
    set_col_widths(sig_tbl, [7.0, 7.0])

    doc.save(output_path)
    print(f"DOCX generated: {output_path}")


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating Professional Construction Bill Documents...")
    build_pdf()
    build_docx()
    print("\nDone! Both files have been generated:")
    print("  - Bill_Final.pdf")
    print("  - Bill_Final.docx")
